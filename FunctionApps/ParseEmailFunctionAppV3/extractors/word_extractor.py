# extractors/word_extractor.py
import io
import logging
import traceback
import zipfile
import re
from xml.etree import ElementTree as ET
# Import shared utility functions
from utils.office_extraction_utils import (
    extract_urls_from_office_html,
    extract_urls_from_relationship_file,
    extract_urls_from_xml_file,
    extract_urls_from_drawing_files,
    filter_microsoft_urls
)

logger = logging.getLogger(__name__)

def extract_text_from_word(word_data):
    """
    Extract text content from Word document binary data.
    
    Args:
        word_data (bytes): The binary content of the Word document
        
    Returns:
        dict: Dictionary containing the extracted text and URLs
    """
    try:
        # Create a file-like object from the binary data
        word_file = io.BytesIO(word_data)
        
        extracted_text = None
        urls = set()
        metadata = []
        
        # Try docx2txt first for simple extraction from docx
        try:
            import docx2txt
            
            # Try reading with docx2txt - works for docx files
            text_content = docx2txt.process(word_file)
            
            if text_content and len(text_content.strip()) > 0:
                logger.debug(f"Successfully extracted Word text using docx2txt, {len(text_content)} characters")
                extracted_text = text_content
            
        except Exception as docx2txt_err:
            logger.warning(f"docx2txt Word reading failed: {str(docx2txt_err)}, trying python-docx")
            
            # Reset file position
            word_file.seek(0)
            
            # Try python-docx
            try:
                import docx
                
                doc = docx.Document(word_file)
                paragraphs = [paragraph.text for paragraph in doc.paragraphs]
                text_content = "\n\n".join(paragraphs)
                
                # Extract table content if available
                for table in doc.tables:
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            row_text.append(cell.text)
                        text_content += "\n" + "\t".join(row_text)
                
                if text_content and len(text_content.strip()) > 0:
                    logger.debug(f"Successfully extracted Word text using python-docx, {len(text_content)} characters")
                    extracted_text = text_content
                    
            except Exception as python_docx_err:
                logger.warning(f"python-docx Word reading failed: {str(python_docx_err)}")
        
        # Always run deep extraction for URLs and metadata, regardless of whether text was extracted
        word_file.seek(0)
        try:
            logger.debug("Running deep extraction for URLs and metadata")
            deep_result = extract_word_deep(word_data)
            
            # Extract URLs from the deep extraction result
            if deep_result:
                # If we got URL data in the deep extraction, extract it
                url_section_match = re.search(r"# URLs\n(.*?)(?:\n\n|$)", deep_result, re.DOTALL)
                if url_section_match:
                    deep_urls = url_section_match.group(1).strip().split("\n")
                    urls.update(deep_urls)
                    logger.debug(f"Extracted {len(deep_urls)} URLs from deep extraction")
                
                # If we didn't get text from simpler methods, use the deep extraction text
                if not extracted_text:
                    extracted_text = deep_result
        except Exception as deep_err:
            logger.warning(f"Deep Word extraction failed: {str(deep_err)}, falling back to antiword")
            
            # If deep extraction fails and we still don't have text, try antiword
            if not extracted_text:
                word_file.seek(0)
                try:
                    import tempfile
                    import subprocess
                    import os
                    
                    # Create a temporary file
                    with tempfile.NamedTemporaryFile(delete=False) as temp_file:
                        temp_file.write(word_data)
                        temp_file_path = temp_file.name
                    
                    try:
                        # Run antiword on the temp file
                        result = subprocess.run(
                            ["antiword", temp_file_path], 
                            capture_output=True, 
                            text=True, 
                            check=True
                        )
                        
                        text_content = result.stdout
                        
                        if text_content and len(text_content.strip()) > 0:
                            logger.debug(f"Successfully extracted Word text using antiword, {len(text_content)} characters")
                            extracted_text = text_content
                    except (FileNotFoundError, subprocess.SubprocessError) as e:
                        logger.warning(f"Antiword execution failed: {str(e)}")
                    finally:
                        # Clean up temp file
                        try:
                            os.unlink(temp_file_path)
                        except Exception as e:
                            logger.warning(f"Failed to remove temporary file: {str(e)}")
                
                except Exception as antiword_err:
                    logger.error(f"All Word reading methods failed: {str(antiword_err)}")
        
        # Convert all URLs to the format expected by url_processor.py
        formatted_urls = []
        for url in urls:
            formatted_urls.append({
                "original_url": url,
                "is_shortened": False,  # This will be processed later
                "expanded_url": ""      # This will be processed later
            })
        
        # If we have URLs, append them to the extracted text
        if urls:
            url_text = "\n\n# URLs\n" + "\n".join(urls)
            logger.debug(f"Adding {len(urls)} URLs to output")
            
            if extracted_text:
                # Only append URLs if they're not already in the text
                if "# URLs" not in extracted_text:
                    extracted_text += url_text
            else:
                # If we somehow have URLs but no text
                extracted_text = "[No text content found in Word document]" + url_text
        
        # Create a result object that includes both text and URLs
        result = {
            "text": extracted_text if extracted_text else "[Error: Could not extract Word text with any available method]",
            "urls": formatted_urls
        }
        
        return result
    
    except Exception as e:
        logger.error(f"Error extracting text from Word document: {str(e)}")
        logger.debug(traceback.format_exc())
        return {
            "text": f"[Error extracting Word text: {str(e)}]",
            "urls": []
        }

def extract_word_deep(word_data):
    """
    Extract detailed information from Word files including text, URLs, VBA code, etc.
    
    Args:
        word_data (bytes): The binary content of the Word file
        
    Returns:
        str: Extracted text content and metadata
    """
    logger.critical("ENTRY: extract_word_deep called")
    word_bytes = io.BytesIO(word_data)
    extracted_text = []
    metadata = []
    urls = set()

    try:
        logger.info("Performing deep extraction from Word file")
        
        with zipfile.ZipFile(word_bytes, 'r') as z:
            # Log all files in the archive
            all_files = z.namelist()
            logger.critical(f"Files in Word document: {all_files}")
            
            # Log specific files we're looking for
            html_files = [f for f in all_files if f.endswith('.htm') or f.endswith('.html')]
            logger.critical(f"HTML files found: {html_files}")
            
            rels_files = [f for f in all_files if f.endswith('.rels')]
            logger.critical(f"Relationship files found: {rels_files}")
            
            drawing_files = [f for f in all_files if 
                            f.endswith('.vml') or 
                            '/drawings/' in f or 
                            '/diagrams/' in f]
            logger.critical(f"Drawing files found: {drawing_files}")
            
            custom_xml_files = [f for f in all_files if 
                               'customXml/' in f and f.endswith('.xml')]
            logger.critical(f"Custom XML files found: {custom_xml_files}")
            
            # Process document.xml which contains the main text
            if 'word/document.xml' in z.namelist():
                try:
                    content = z.read('word/document.xml').decode('utf-8', errors='ignore')
                    root = ET.fromstring(content)
                    
                    # Extract all text from paragraph (w:p) and text (w:t) elements
                    for elem in root.iter():
                        if elem.tag.endswith('}t'):  # Text element in Word XML
                            if elem.text and elem.text.strip():
                                extracted_text.append(elem.text.strip())
                    
                    # Extract URLs from document.xml using regex
                    logger.critical("About to call extract_urls_from_xml_file for document.xml")
                    xml_urls = extract_urls_from_xml_file(z, 'word/document.xml')
                    logger.critical(f"extract_urls_from_xml_file returned {len(xml_urls)} URLs")
                    urls.update(xml_urls)
                    
                except ET.ParseError as e:
                    logger.warning(f"Could not parse document.xml: {str(e)}")
                except Exception as e:
                    logger.warning(f"Error processing document.xml: {str(e)}")
            
            # Process headers
            for filename in z.namelist():
                if filename.startswith('word/header'):
                    try:
                        content = z.read(filename).decode('utf-8', errors='ignore')
                        root = ET.fromstring(content)
                        
                        header_text = []
                        for elem in root.iter():
                            if elem.tag.endswith('}t'):
                                if elem.text and elem.text.strip():
                                    header_text.append(elem.text.strip())
                        
                        if header_text:
                            metadata.append("# Header")
                            metadata.append(" ".join(header_text))
                        
                        # Extract URLs from header files
                        logger.critical(f"About to call extract_urls_from_xml_file for {filename}")
                        xml_urls = extract_urls_from_xml_file(z, filename)
                        logger.critical(f"extract_urls_from_xml_file returned {len(xml_urls)} URLs")
                        urls.update(xml_urls)
                        
                    except Exception as e:
                        logger.warning(f"Error processing header {filename}: {str(e)}")
            
            # Process footers
            for filename in z.namelist():
                if filename.startswith('word/footer'):
                    try:
                        content = z.read(filename).decode('utf-8', errors='ignore')
                        root = ET.fromstring(content)
                        
                        footer_text = []
                        for elem in root.iter():
                            if elem.tag.endswith('}t'):
                                if elem.text and elem.text.strip():
                                    footer_text.append(elem.text.strip())
                        
                        if footer_text:
                            metadata.append("# Footer")
                            metadata.append(" ".join(footer_text))
                        
                        # Extract URLs from footer files
                        logger.critical(f"About to call extract_urls_from_xml_file for {filename}")
                        xml_urls = extract_urls_from_xml_file(z, filename)
                        logger.critical(f"extract_urls_from_xml_file returned {len(xml_urls)} URLs")
                        urls.update(xml_urls)
                        
                    except Exception as e:
                        logger.warning(f"Error processing footer {filename}: {str(e)}")
            
            # Process comments
            if 'word/comments.xml' in z.namelist():
                try:
                    content = z.read('word/comments.xml').decode('utf-8', errors='ignore')
                    root = ET.fromstring(content)
                    
                    comments = []
                    for elem in root.iter():
                        if elem.tag.endswith('}comment'):
                            comment_text = []
                            for text_elem in elem.iter():
                                if text_elem.tag.endswith('}t') and text_elem.text:
                                    comment_text.append(text_elem.text)
                            if comment_text:
                                comments.append(" ".join(comment_text))
                    
                    if comments:
                        metadata.append("# Comments")
                        metadata.append("\n".join(comments))
                    
                    # Extract URLs from comments
                    logger.critical(f"About to call extract_urls_from_xml_file for word/comments.xml")
                    xml_urls = extract_urls_from_xml_file(z, 'word/comments.xml')
                    logger.critical(f"extract_urls_from_xml_file returned {len(xml_urls)} URLs")
                    urls.update(xml_urls)
                    
                except Exception as e:
                    logger.warning(f"Error processing comments.xml: {str(e)}")
            
            # Process VBA code if present
            if 'word/vbaProject.bin' in z.namelist():
                try:
                    vba_binary = z.read('word/vbaProject.bin')
                    # Look for VBA code segments
                    modules = re.findall(b'Attribute VB_Name = "([^"]+)".*?(?=Attribute VB_Name|$)', 
                                       vba_binary, re.DOTALL)
                    
                    if modules:
                        metadata.append("# Word VBA Code")
                    
                    for i, module in enumerate(modules):
                        try:
                            decoded_content = module.decode('utf-8', errors='ignore')
                            cleaned_content = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f]', '', decoded_content)
                            if cleaned_content.strip():
                                metadata.append(f"# Module {i}")
                                metadata.append(cleaned_content)
                                
                                # Look for URLs in VBA code
                                vba_urls = re.findall(r'https?://(?:[-\w.]|(?:%[\da-fA-F]{2}))+[^\s"\'<>]*', cleaned_content)
                                urls.update(vba_urls)
                                
                        except Exception as e:
                            logger.warning(f"Could not decode VBA module {i}: {str(e)}")
                except Exception as e:
                    logger.warning(f"Error processing VBA project: {str(e)}")
            
            # Process relationship files for hyperlinks
            for rels_file in rels_files:
                logger.critical(f"About to call extract_urls_from_relationship_file for {rels_file}")
                rel_urls = extract_urls_from_relationship_file(z, rels_file)
                logger.critical(f"extract_urls_from_relationship_file returned {len(rel_urls)} URLs")
                urls.update(rel_urls)
            
            # Process HTML files, especially afchunk.htm which contains image links
            for html_file in html_files:
                logger.critical(f"About to call extract_urls_from_office_html for {html_file}")
                html_urls = extract_urls_from_office_html(z, html_file)
                logger.critical(f"extract_urls_from_office_html returned {len(html_urls)} URLs")
                urls.update(html_urls)
            
            # Check drawing files (often contain linked images)
            for drawing_file in drawing_files:
                logger.critical(f"About to call extract_urls_from_drawing_files for {drawing_file}")
                drawing_urls = extract_urls_from_drawing_files(z, drawing_file)
                logger.critical(f"extract_urls_from_drawing_files returned {len(drawing_urls)} URLs")
                urls.update(drawing_urls)
            
            # Check custom XML parts
            for xml_file in custom_xml_files:
                logger.critical(f"About to call extract_urls_from_xml_file for {xml_file}")
                xml_urls = extract_urls_from_xml_file(z, xml_file)
                logger.critical(f"extract_urls_from_xml_file returned {len(xml_urls)} URLs")
                urls.update(xml_urls)
            
            # Filter URLs to remove Microsoft-related ones
            logger.critical(f"About to call filter_microsoft_urls with {len(urls)} URLs")
            urls = filter_microsoft_urls(urls)
            logger.critical(f"filter_microsoft_urls returned {len(urls)} URLs")
            
            if urls:
                metadata.append("# URLs")
                metadata.append("\n".join(urls))
            
            # Check for embedded objects
            embedded_files = []
            for filename in z.namelist():
                if 'word/embeddings' in filename:
                    try:
                        embedded_files.append(f"{filename} ({z.getinfo(filename).file_size} bytes)")
                    except Exception as e:
                        logger.warning(f"Error processing embedded file {filename}: {str(e)}")
            
            if embedded_files:
                metadata.append("# Embedded Files")
                metadata.append("\n".join(embedded_files))

        # Clean and join the extracted text
        if extracted_text:
            full_text = ' '.join(extracted_text)
            cleaned_text = re.sub(r'[^\x20-\x7E\n\r]+', '', full_text)
            cleaned_text = re.sub(r'\s+', ' ', cleaned_text).strip()
            
            # Combine main text and metadata
            result = cleaned_text
            if metadata:
                result += "\n\n" + "\n\n".join(metadata)
            
            logger.critical("EXIT: extract_word_deep with success")
            return result
        elif metadata:
            # Return just metadata if no text content
            logger.critical("EXIT: extract_word_deep with metadata only")
            return "\n\n".join(metadata)
        else:
            logger.critical("EXIT: extract_word_deep with no content")
            return "[No text content found in Word document]"

    except zipfile.BadZipFile:
        logger.warning("Not a valid DOCX file (bad zip format), passing to other extractors")
        raise
    except Exception as e:
        logger.error(f"Failed to extract text from Word document (deep extraction): {str(e)}")
        logger.debug(traceback.format_exc())
        raise